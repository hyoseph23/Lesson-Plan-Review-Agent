"""
K-12 Lesson Plan Review Agent — Web Interface
"""
import os, json, tempfile, io
from pathlib import Path
from flask import Flask, request, render_template_string, jsonify, send_file
from agent import review_document

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>RainCheck: Lesson Plan Generation & Review Agent</title>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,wght@0,300;0,500;0,700;1,300&family=DM+Mono:wght@400;500&family=Nunito:wght@400;500;600;700&display=swap" rel="stylesheet"/>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg:#fffdf8;
  --surface:#ffffff;
  --surface2:#faf8f3;
  --border:#e8e2d6;
  --border2:#d4ccba;
  --orange:#e8651a;
  --orange-light:#fff3eb;
  --orange-mid:#fde4cc;
  --green:#2d7a4f;
  --green-light:#edf7f2;
  --green-mid:#c4e8d4;
  --red:#c0392b;
  --red-light:#fdf0ef;
  --text:#1a1a14;
  --text2:#5c5748;
  --text3:#9c9485;
  --shadow:0 1px 3px rgba(0,0,0,.06), 0 4px 16px rgba(0,0,0,.04);
  --shadow-lg:0 2px 8px rgba(0,0,0,.08), 0 12px 40px rgba(0,0,0,.06);
}
body{font-family:'Nunito',sans-serif;background:var(--bg);color:var(--text);min-height:100vh;padding:0}

/* ── Header ── */
.topbar{
  background:var(--surface);border-bottom:1.5px solid var(--border);
  padding:20px 48px;display:flex;align-items:center;justify-content:space-between;
  position:sticky;top:0;z-index:100;
}
.brand{display:flex;align-items:center;gap:12px}
.pencil-icon{
  width:38px;height:38px;background:var(--orange);border-radius:10px;
  display:flex;align-items:center;justify-content:center;font-size:18px;
  box-shadow:0 2px 8px rgba(232,101,26,.3);flex-shrink:0;
}
.brand-text{}
.brand-name{font-family:'Fraunces',serif;font-size:18px;font-weight:700;color:var(--text);line-height:1}
.brand-sub{font-size:11px;color:var(--text3);font-family:'DM Mono',monospace;letter-spacing:.08em;margin-top:2px}
.badge{background:var(--orange-light);border:1px solid var(--orange-mid);color:var(--orange);font-family:'DM Mono',monospace;font-size:10px;letter-spacing:.1em;padding:4px 10px;border-radius:20px}

/* ── Tabs ── */
.tabs{display:flex;gap:0;border-bottom:1.5px solid var(--border);padding:0 48px;background:var(--surface)}
.tab{padding:14px 24px;font-size:14px;font-weight:600;cursor:pointer;border-bottom:2.5px solid transparent;
  color:var(--text3);transition:all .15s;margin-bottom:-1.5px;font-family:'Nunito',sans-serif}
.tab:hover{color:var(--text2)}
.tab.active{color:var(--orange);border-bottom-color:var(--orange)}

/* ── Main layout ── */
.main{max-width:1140px;margin:0 auto;padding:48px 48px 120px}
.tab-page{display:none}
.tab-page.active{display:block}

/* ── Hero ── */
.hero{margin-bottom:48px}
.hero h1{font-family:'Fraunces',serif;font-size:42px;font-weight:700;line-height:1.1;color:var(--text);margin-bottom:12px}
.hero h1 em{font-style:italic;color:var(--orange)}
.hero p{font-size:15px;color:var(--text2);max-width:520px;line-height:1.6}

/* ── Upload card ── */
.card{background:var(--surface);border:1.5px solid var(--border);border-radius:16px;padding:32px;box-shadow:var(--shadow);margin-bottom:24px}
.card-title{font-family:'Fraunces',serif;font-size:16px;font-weight:500;color:var(--text);margin-bottom:20px;display:flex;align-items:center;gap:8px}
.card-title span{color:var(--orange)}
label.file-label{display:block;font-size:13px;color:var(--text2);margin-bottom:8px;font-weight:600}
input[type=file]{
  display:block;width:100%;font-size:13px;color:var(--text2);
  background:var(--surface2);border:1.5px solid var(--border);border-radius:8px;
  padding:10px 14px;margin-bottom:20px;cursor:pointer;font-family:'Nunito',sans-serif;
}
input[type=file]::file-selector-button{
  background:var(--orange);color:#fff;border:none;padding:6px 16px;
  border-radius:6px;font-size:12px;font-weight:700;cursor:pointer;margin-right:14px;
  font-family:'Nunito',sans-serif;letter-spacing:.02em;
}
.run-btn{
  background:var(--orange);color:#fff;border:none;padding:13px 32px;
  border-radius:10px;font-family:'Nunito',sans-serif;font-size:14px;font-weight:700;
  cursor:pointer;letter-spacing:.02em;display:inline-flex;align-items:center;gap:8px;
  box-shadow:0 2px 12px rgba(232,101,26,.35);transition:all .15s;
}
.run-btn:hover{background:#d4570f;transform:translateY(-1px);box-shadow:0 4px 16px rgba(232,101,26,.4)}
.run-btn:disabled{opacity:.45;cursor:not-allowed;transform:none}
.status{font-family:'DM Mono',monospace;font-size:12px;color:var(--orange);margin-top:14px;min-height:18px}
.error-box{color:var(--red);background:var(--red-light);border:1px solid #f5c6c3;padding:12px 16px;border-radius:8px;font-size:13px;margin-top:14px;display:none}

/* ── Summary bar ── */
.summary-bar{
  background:var(--surface);border:1.5px solid var(--border);border-radius:12px;
  padding:16px 24px;display:flex;align-items:center;gap:28px;
  margin-bottom:24px;flex-wrap:wrap;box-shadow:var(--shadow);
}
.summary-label{font-family:'DM Mono',monospace;font-size:10px;color:var(--text3);letter-spacing:.12em;text-transform:uppercase}
.summary-item{display:flex;align-items:center;gap:8px}
.summary-icon{font-size:15px}
.summary-text{font-size:13px;color:var(--text2)}
.summary-count{font-family:'DM Mono',monospace;font-size:15px;font-weight:700}
.summary-divider{width:1px;height:24px;background:var(--border)}

/* ── Results header ── */
#results{display:none}
.results-hdr{display:flex;align-items:center;justify-content:space-between;margin-bottom:20px;flex-wrap:wrap;gap:12px}
.results-title{font-family:'Fraunces',serif;font-size:22px;font-weight:700;color:var(--text)}
.hdr-actions{display:flex;gap:8px;align-items:center}
.issue-pill{font-family:'DM Mono',monospace;font-size:11px;padding:5px 12px;border-radius:20px}
.issue-pill.has{background:var(--red-light);color:var(--red);border:1px solid #f5c6c3}
.issue-pill.none{background:var(--green-light);color:var(--green);border:1px solid var(--green-mid)}
.dl-btn{background:var(--surface);border:1.5px solid var(--border);color:var(--text2);padding:7px 16px;border-radius:8px;font-family:'DM Mono',monospace;font-size:11px;cursor:pointer;transition:all .15s}
.dl-btn:hover{border-color:var(--orange);color:var(--orange)}

/* ── Review table ── */
.review-table-wrap{border:1.5px solid var(--border);border-radius:12px;overflow:hidden;box-shadow:var(--shadow);margin-bottom:32px}
.tbl-header{
  display:grid;grid-template-columns:200px 1fr 1fr 160px;
  background:var(--surface2);border-bottom:1.5px solid var(--border);
  padding:12px 16px;gap:12px;
}
.tbl-header-cell{font-family:'DM Mono',monospace;font-size:10px;letter-spacing:.12em;text-transform:uppercase;color:var(--text3);font-weight:500}
.tbl-row{
  display:grid;grid-template-columns:200px 1fr 1fr 160px;
  border-bottom:1px solid var(--border);padding:16px;gap:12px;
  background:var(--surface);transition:background .1s;align-items:start;
}
.tbl-row:last-child{border-bottom:none}
.tbl-row:hover{background:var(--surface2)}
.section-name{font-family:'DM Mono',monospace;font-size:11px;font-weight:500;color:var(--text2);line-height:1.5}
.bullets{list-style:none;display:flex;flex-direction:column;gap:5px}
.bullets li{display:flex;gap:7px;align-items:flex-start;font-size:12.5px;line-height:1.5;color:var(--text2)}
.bullets li .bicon{flex-shrink:0;margin-top:2px;font-size:11px}
.ok-item .bicon{color:var(--green)}
.err-item{color:var(--red)!important}
.err-item .bicon{color:var(--red)}
.warn-item{color:#b85c00!important}
.warn-item .bicon{color:var(--orange)}

/* ── Accept/Override controls ── */
.action-cell{display:flex;flex-direction:column;gap:6px}
.accept-btn,.override-btn{
  width:100%;padding:7px 10px;border-radius:7px;font-size:11px;font-weight:700;
  cursor:pointer;font-family:'Nunito',sans-serif;letter-spacing:.02em;
  display:flex;align-items:center;justify-content:center;gap:5px;transition:all .15s;border:1.5px solid;
}
.accept-btn{background:var(--green-light);color:var(--green);border-color:var(--green-mid)}
.accept-btn:hover{background:var(--green-mid)}
.accept-btn.active{background:var(--green);color:#fff;border-color:var(--green)}
.override-btn{background:var(--orange-light);color:var(--orange);border-color:var(--orange-mid)}
.override-btn:hover{background:var(--orange-mid)}
.override-btn.active{background:var(--orange);color:#fff;border-color:var(--orange)}
.override-input{
  display:none;width:100%;border:1.5px solid var(--orange-mid);border-radius:6px;
  padding:6px 8px;font-size:11px;font-family:'Nunito',sans-serif;color:var(--text);
  background:var(--surface);margin-top:4px;resize:vertical;min-height:60px;
}
.override-input.show{display:block}

/* ── Question form ── */
.question-form{display:flex;flex-direction:column;gap:20px}
.q-item{display:flex;flex-direction:column;gap:6px}
.q-label{font-size:13px;font-weight:700;color:var(--text)}
.q-hint{font-size:12px;color:var(--text3);margin-top:-4px}
.q-input{background:var(--surface2);border:1.5px solid var(--border);border-radius:8px;
  padding:10px 14px;font-size:13px;font-family:'Nunito',sans-serif;color:var(--text);
  transition:border .15s;width:100%}
.q-input:focus{outline:none;border-color:var(--orange)}
textarea.q-input{resize:vertical;min-height:72px}
.q-row{display:grid;grid-template-columns:1fr 1fr;gap:16px}
.create-btn{background:var(--green);color:#fff;border:none;padding:13px 32px;
  border-radius:10px;font-family:'Nunito',sans-serif;font-size:14px;font-weight:700;
  cursor:pointer;display:inline-flex;align-items:center;gap:8px;
  box-shadow:0 2px 12px rgba(45,122,79,.3);transition:all .15s;margin-top:8px}
.create-btn:hover{background:#256b44;transform:translateY(-1px)}
.create-btn:disabled{opacity:.45;cursor:not-allowed;transform:none}
.create-status{font-family:'DM Mono',monospace;font-size:12px;color:var(--green);margin-top:12px;min-height:18px}
.progress-steps{display:flex;flex-direction:column;gap:8px;margin-top:16px}
.progress-step{display:flex;align-items:center;gap:10px;font-size:13px;color:var(--text3)}
.progress-step.active{color:var(--orange);font-weight:600}
.progress-step.done{color:var(--green)}
.step-dot{width:8px;height:8px;border-radius:50%;background:var(--border);flex-shrink:0}
.progress-step.active .step-dot{background:var(--orange)}
.progress-step.done .step-dot{background:var(--green)}

/* ── Generate button ── */
.generate-bar{
  position:fixed;bottom:0;left:0;right:0;
  background:rgba(255,253,248,.92);backdrop-filter:blur(12px);
  border-top:1.5px solid var(--border);padding:16px 48px;
  display:none;align-items:center;justify-content:space-between;gap:16px;z-index:90;
}
.generate-bar.show{display:flex}
.generate-hint{font-size:13px;color:var(--text2)}
.generate-hint strong{color:var(--text)}
.generate-btn{
  background:var(--green);color:#fff;border:none;padding:14px 36px;
  border-radius:10px;font-family:'Nunito',sans-serif;font-size:14px;font-weight:700;
  cursor:pointer;display:inline-flex;align-items:center;gap:10px;
  box-shadow:0 2px 12px rgba(45,122,79,.35);transition:all .15s;letter-spacing:.02em;
}
.generate-btn:hover{background:#256b44;transform:translateY(-1px);box-shadow:0 4px 16px rgba(45,122,79,.4)}
.generate-btn:disabled{opacity:.45;cursor:not-allowed;transform:none}
</style>
</head>
<body>

<!-- Top bar -->
<div class="topbar">
  <div class="brand">
    <div class="pencil-icon">✏️</div>
    <div class="brand-text">
      <div class="brand-name">RainCheck</div>
      <div class="brand-sub">K–12 Lesson Review</div>
    </div>
  </div>
  <div class="badge">AI-Powered · v1.0</div>
</div>

<!-- Tabs -->
<div class="tabs">
  <div class="tab active" id="tab-create" onclick="switchTab('create')">Create Plan</div>
  <div class="tab" id="tab-review" onclick="switchTab('review')">Review Plan</div>
</div>

<div class="main">

  <!-- CREATE TAB -->
  <div class="tab-page active" id="page-create">
    <div class="hero">
      <h1>Create a plan<br/>for your <em>substitute</em></h1>
      <p>Answer a few quick questions and get a complete, formatted lesson plan ready to hand to your sub before you leave.</p>
    </div>
    <div class="card">
      <div class="card-title">✏️ <span>Tell us about your class</span></div>
      <div class="question-form">
        <div class="q-row">
          <div class="q-item">
            <label class="q-label" for="q-grade">What grade are you teaching?</label>
            <input class="q-input" id="q-grade" type="text" placeholder="e.g. Grade 4"/>
          </div>
          <div class="q-item">
            <label class="q-label" for="q-subject">What subject is this for?</label>
            <input class="q-input" id="q-subject" type="text" placeholder="e.g. Science, Math, ELA"/>
          </div>
        </div>
        <div class="q-row">
          <div class="q-item">
            <label class="q-label" for="q-duration">How long is each class period?</label>
            <input class="q-input" id="q-duration" type="text" placeholder="e.g. 45 min"/>
          </div>
          <div class="q-item">
            <label class="q-label" for="q-days">How many days will you be out?</label>
            <input class="q-input" id="q-days" type="text" placeholder="e.g. 2 days, just Friday"/>
          </div>
        </div>
        <div class="q-item">
          <label class="q-label" for="q-unit">What unit are you currently in, and what did students learn last week?</label>
          <div class="q-hint">Give the sub context so they understand where the class is in the curriculum.</div>
          <textarea class="q-input" id="q-unit" placeholder="e.g. We are in Unit 3: Fractions and Decimals. Last week students learned how to add and subtract fractions with like denominators. They completed the Unit 3 quiz on Friday and did well overall."></textarea>
        </div>
        <div class="q-item">
          <label class="q-label" for="q-topic">What should the sub cover during your absence?</label>
          <div class="q-hint">Be as specific as you like — reference page numbers, activities, or learning goals.</div>
          <textarea class="q-input" id="q-topic" placeholder="e.g. The sub should introduce decimals using textbook pages 42-45 and the practice worksheet in the green folder on my desk."></textarea>
        </div>
        <div class="q-item">
          <label class="q-label" for="q-level">How much should the sub manage vs just facilitate?</label>
          <div class="q-hint">This helps set expectations for the sub's involvement level.</div>
          <input class="q-input" id="q-level" type="text" placeholder="e.g. Mostly independent work, sub just needs to keep order. OR Sub should lead a short discussion."/>
        </div>
        <div class="q-item">
          <label class="q-label" for="q-materials">Where can the sub find materials?</label>
          <div class="q-hint">Worksheets, textbooks, digital resources, supply locations.</div>
          <textarea class="q-input" id="q-materials" placeholder="e.g. Worksheets are printed and in the red tray on my desk. Textbooks are on the shelf by the window. The class iPad cart is in the hallway."></textarea>
        </div>
        <div class="q-item">
          <label class="q-label" for="q-notes">Anything else the sub needs to know?</label>
          <div class="q-hint">Class routines, students who need extra support, behavior notes, emergency contacts.</div>
          <textarea class="q-input" id="q-notes" placeholder="e.g. Marcus sits in front and needs reminders to stay on task. The class does independent reading for the first 5 minutes. Lunch is at 11:30."></textarea>
        </div>
      </div>
      <button class="create-btn" id="createbtn" onclick="createPlan()">Generate Sub Plan</button>
      <div class="create-status" id="create-status"></div>
      <div class="progress-steps" id="create-steps" style="display:none">
        <div class="progress-step" id="cstep-1"><div class="step-dot"></div>Reading your answers</div>
        <div class="progress-step" id="cstep-2"><div class="step-dot"></div>Writing lesson plan</div>
        <div class="progress-step" id="cstep-3"><div class="step-dot"></div>Formatting document</div>
      </div>
      <div class="error-box" id="create-errmsg"></div>
    </div>
  </div>

  <!-- REVIEW TAB -->
  <div class="tab-page" id="page-review">
    <div class="hero">
      <h1>Review your<br/><em>lesson plan</em> instantly</h1>
      <p>Upload a lesson plan or worksheet and get section-by-section feedback on formatting, content quality, and standards alignment — in under a minute.</p>
    </div>
    <div class="card">
      <div class="card-title">✏️ <span>Upload Document</span></div>
      <label class="file-label" for="docfile">Lesson plan or worksheet (.docx, .pdf, or .txt)</label>
      <input type="file" id="docfile" name="file" accept=".docx,.pdf,.txt"/>
      <button class="run-btn" id="runbtn" onclick="runReview()">▶&nbsp; Run Review</button>
      <div class="status" id="status"></div>
      <div class="error-box" id="errmsg"></div>
    </div>

  <!-- Results -->
  <div id="results">
    <div class="results-hdr">
      <div class="results-title">Review Results</div>
      <div class="hdr-actions">
        <span class="issue-pill" id="pill"></span>
        <button class="dl-btn" onclick="dlCSV()">↓ CSV</button>
        <button class="dl-btn" onclick="dlWord()">↓ Word</button>
      </div>
    </div>

    <!-- Summary bar -->
    <div class="summary-bar" id="summary-bar"></div>

    <!-- Table -->
    <div class="review-table-wrap">
      <div class="tbl-header">
        <div class="tbl-header-cell">Section</div>
        <div class="tbl-header-cell">Formatting &amp; Grammar</div>
        <div class="tbl-header-cell">Section Feedback</div>
        <div class="tbl-header-cell">Action</div>
      </div>
      <div id="tbl-body"></div>
    </div>
  </div>

</div>
  </div><!-- end page-review -->

</div><!-- end main -->

<!-- Generate bar (fixed bottom) -->
<div class="generate-bar" id="generate-bar">
  <div class="generate-hint">
    <strong id="action-summary">0 sections accepted</strong> — ready to generate your formatted lesson plan
  </div>
  <button class="generate-btn" id="generate-btn" onclick="generatePlan()">
    ✨&nbsp; Generate Formatted Lesson Plan
  </button>
</div>

<script>
let reviewData = null;
// rowDecisions[i] = 'accept' | 'override' | null
let rowDecisions = [];

function classify(t){
  const s = t.toLowerCase();
  if(s==='no issues found') return 'ok-item';
  if(s.includes('missing')||s.includes('not found')||s.includes('absent')) return 'err-item';
  if(s.includes('flag')||s.includes('exceeds')||s.includes('weak')||s.includes('vague')||
     s.includes('unclear')||s.includes('should')||s.includes('must')) return 'warn-item';
  return '';
}

function rowStatus(row){
  const bullets = row.section_feedback || ['No issues found'];
  if(bullets.some(b => classify(b)==='err-item')) return 'red';
  if(bullets.some(b => classify(b)==='warn-item'||(classify(b)===''&&b.toLowerCase()!=='no issues found'))) return 'orange';
  return 'green';
}

function makeBullets(arr){
  const icons = {'ok-item':'✓','err-item':'✗','warn-item':'⚠','':'→'};
  return '<ul class="bullets">'+(arr||['No issues found']).map(b=>{
    const cls = classify(b);
    return `<li class="${cls}"><span class="bicon">${icons[cls]||'→'}</span><span>${b}</span></li>`;
  }).join('')+'</ul>';
}

function renderSummary(rows){
  let counts={green:0,orange:0,red:0};
  rows.forEach(r=>counts[rowStatus(r)]++);
  document.getElementById('summary-bar').innerHTML =
    `<span class="summary-label">Section Status</span>` +
    `<div class="summary-divider"></div>` +
    [
      {icon:'✓',label:'Ready to go',color:'var(--green)',count:counts.green},
      {icon:'⚠',label:'Needs fine-tuning',color:'var(--orange)',count:counts.orange},
      {icon:'✗',label:'Failed review',color:'var(--red)',count:counts.red},
    ].map((item,i)=>
      `${i>0?'<div class="summary-divider"></div>':''}
      <div class="summary-item">
        <span class="summary-icon" style="color:${item.color}">${item.icon}</span>
        <span class="summary-text">${item.label}</span>
        <span class="summary-count" style="color:${item.color}">${item.count}</span>
      </div>`
    ).join('');
}

function updateActionSummary(){
  const accepted = rowDecisions.filter(d=>d==='accept').length;
  const overridden = rowDecisions.filter(d=>d==='override').length;
  const total = accepted + overridden;
  let text = '';
  if(total===0) text = 'No sections actioned yet';
  else if(overridden===0) text = `${accepted} section${accepted!==1?'s':''} accepted`;
  else if(accepted===0) text = `${overridden} section${overridden!==1?'s':''} overridden`;
  else text = `${accepted} accepted, ${overridden} overridden`;
  document.getElementById('action-summary').textContent = text;
  document.getElementById('generate-btn').disabled = total === 0;
}

function setDecision(idx, decision){
  const prev = rowDecisions[idx];

  // Toggle off if same button clicked again
  if(prev === decision){
    rowDecisions[idx] = null;
    document.getElementById(`accept-${idx}`).classList.remove('active');
    document.getElementById(`override-${idx}`).classList.remove('active');
    const oi = document.getElementById(`override-input-${idx}`);
    if(oi){ oi.classList.remove('show'); }
  } else {
    rowDecisions[idx] = decision;
    document.getElementById(`accept-${idx}`).classList.toggle('active', decision==='accept');
    document.getElementById(`override-${idx}`).classList.toggle('active', decision==='override');
    const oi = document.getElementById(`override-input-${idx}`);
    if(oi){ oi.classList.toggle('show', decision==='override'); }
  }
  updateActionSummary();
}

function renderTable(rows){
  rowDecisions = new Array(rows.length).fill(null);
  let issues=0;

  document.getElementById('tbl-body').innerHTML = rows.map((r,idx)=>{
    const fb = r.formatting_feedback||['No issues found'];
    const sb = r.section_feedback||['No issues found'];
    [...fb,...sb].forEach(b=>{ if(b.toLowerCase()!=='no issues found') issues++; });
    const allOk = rowStatus(r)==='green';

    return `<div class="tbl-row" id="row-${idx}">
      <div class="section-name">${r.section}</div>
      <div>${makeBullets(fb)}</div>
      <div>${makeBullets(sb)}</div>
      <div class="action-cell">
        <button class="accept-btn${allOk?' active':''}" id="accept-${idx}"
          onclick="setDecision(${idx},'accept')">
          ✓ Accept
        </button>
        <button class="override-btn" id="override-${idx}"
          onclick="setDecision(${idx},'override')">
          ✏ Override
        </button>
        <textarea class="override-input" id="override-input-${idx}"
          placeholder="Describe your override or paste corrected text…"></textarea>
      </div>
    </div>`;
  }).join('');

  // Auto-accept sections with no issues
  rows.forEach((r,idx)=>{ if(rowStatus(r)==='green') rowDecisions[idx]='accept'; });

  const pill = document.getElementById('pill');
  pill.textContent = issues>0 ? issues+' issue(s) found' : 'No issues found';
  pill.className = 'issue-pill '+(issues>0?'has':'none');

  renderSummary(rows);
  updateActionSummary();
  document.getElementById('generate-bar').classList.add('show');
}

async function runReview(){
  const fi = document.getElementById('docfile');
  if(!fi.files||!fi.files[0]){ alert('Please select a file first.'); return; }
  const file = fi.files[0];
  const ext = file.name.split('.').pop().toLowerCase();
  if(!['docx','pdf','txt'].includes(ext)){ alert('Please upload a .docx, .pdf, or .txt file.'); return; }

  document.getElementById('errmsg').style.display='none';
  document.getElementById('results').style.display='none';
  document.getElementById('generate-bar').classList.remove('show');
  document.getElementById('runbtn').disabled=true;
  document.getElementById('status').textContent='Reviewing your lesson plan… ~30 seconds';

  const fd = new FormData();
  fd.append('file', file);

  try{
    const res = await fetch('/review', {method:'POST', body:fd});
    const data = await res.json();
    if(!res.ok||data.error) throw new Error(data.error||'Server error');
    reviewData = data.rows;
    renderTable(reviewData);
    document.getElementById('status').textContent='';
    document.getElementById('results').style.display='block';
    document.getElementById('results').scrollIntoView({behavior:'smooth',block:'start'});
  } catch(e){
    document.getElementById('status').textContent='';
    const em = document.getElementById('errmsg');
    em.textContent='Error: '+e.message;
    em.style.display='block';
  }
  document.getElementById('runbtn').disabled=false;
}

async function generatePlan(){
  if(!reviewData) return;
  const decisions = reviewData.map((r,i)=>({
    section: r.section,
    decision: rowDecisions[i]||'pending',
    override_note: document.getElementById(`override-input-${i}`)?.value||'',
    formatting_feedback: r.formatting_feedback,
    section_feedback: r.section_feedback,
  }));

  document.getElementById('generate-btn').disabled=true;
  document.getElementById('generate-btn').textContent='Generating…';

  try{
    const res = await fetch('/generate', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({rows: reviewData, decisions})
    });
    if(!res.ok){ alert('Generation failed.'); return; }
    const blob = await res.blob();
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = 'formatted-lesson-plan.docx';
    a.click();
  } catch(e){
    alert('Error: '+e.message);
  }

  document.getElementById('generate-btn').disabled=false;
  document.getElementById('generate-btn').innerHTML='✨&nbsp; Generate Formatted Lesson Plan';
}

function dlCSV(){
  if(!reviewData) return;
  const rows=[['Section','Formatting Feedback','Section Feedback','Decision','Override Note'],
    ...reviewData.map((r,i)=>[
      r.section,
      (r.formatting_feedback||[]).join(' | '),
      (r.section_feedback||[]).join(' | '),
      rowDecisions[i]||'pending',
      document.getElementById(`override-input-${i}`)?.value||''
    ])];
  const csv=rows.map(r=>r.map(c=>'"'+String(c).replace(/"/g,'""')+'"').join(',')).join('\\n');
  const a=document.createElement('a');
  a.href=URL.createObjectURL(new Blob([csv],{type:'text/csv'}));
  a.download='lesson-review.csv'; a.click();
}

async function dlWord(){
  if(!reviewData) return;
  const decisions = reviewData.map((r,i)=>({
    section:r.section, decision:rowDecisions[i]||'pending',
    override_note: document.getElementById(`override-input-${i}`)?.value||'',
    formatting_feedback:r.formatting_feedback, section_feedback:r.section_feedback
  }));
  const res = await fetch('/download-docx', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({rows: reviewData, decisions})
  });
  if(!res.ok){ alert('Failed to generate Word document.'); return; }
  const blob = await res.blob();
  const a=document.createElement('a'); a.href=URL.createObjectURL(blob);
  a.download='lesson-review.docx'; a.click();
}

function switchTab(tab){
  document.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));
  document.querySelectorAll('.tab-page').forEach(p=>p.classList.remove('active'));
  document.getElementById('tab-'+tab).classList.add('active');
  document.getElementById('page-'+tab).classList.add('active');
  if(tab!=='review') document.getElementById('generate-bar').style.display='none';
}

function setCStep(n){
  for(let i=1;i<=3;i++){
    const el=document.getElementById('cstep-'+i);
    if(!el) continue;
    el.classList.remove('active','done');
    if(i<n) el.classList.add('done');
    if(i===n) el.classList.add('active');
  }
}

async function createPlan(){
  const grade    = document.getElementById('q-grade').value.trim();
  const subject  = document.getElementById('q-subject').value.trim();
  const duration = document.getElementById('q-duration').value.trim();
  const days     = document.getElementById('q-days').value.trim();
  const unit     = document.getElementById('q-unit').value.trim();
  const topic    = document.getElementById('q-topic').value.trim();
  const level    = document.getElementById('q-level').value.trim();
  const materials= document.getElementById('q-materials').value.trim();
  const notes    = document.getElementById('q-notes').value.trim();

  if(!grade||!subject||!unit||!topic){
    alert('Please fill in Grade, Subject, and Topic before generating.');
    return;
  }

  document.getElementById('create-errmsg').style.display='none';
  document.getElementById('createbtn').disabled=true;
  document.getElementById('create-status').textContent='';
  document.getElementById('create-steps').style.display='flex';
  setCStep(1);

  try{
    setCStep(2);
    const res = await fetch('/create', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({grade,subject,duration,days,unit,topic,level,materials,notes})
    });
    setCStep(3);
    if(!res.ok){ const e=await res.json().catch(()=>({})); throw new Error(e.error||'Server error'); }
    const blob = await res.blob();
    const a=document.createElement('a'); a.href=URL.createObjectURL(blob);
    a.download='sub-lesson-plan.docx'; a.click();
    document.getElementById('create-status').textContent='Plan downloaded successfully.';
    document.getElementById('create-steps').style.display='none';
  } catch(e){
    document.getElementById('create-steps').style.display='none';
    const em=document.getElementById('create-errmsg');
    em.textContent='Error: '+e.message; em.style.display='block';
  }
  document.getElementById('createbtn').disabled=false;
}
</script>
</body>
</html>"""


@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/review", methods=["POST"])
def review():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("docx", "pdf", "txt"):
        return jsonify({"error": "Only .docx, .pdf, and .txt files are supported"}), 400
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name
    try:
        rows = review_document(tmp_path)
        return jsonify({"rows": rows})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def build_review_docx(rows, decisions=None):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    decisions = decisions or []
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    run = title.add_run("✏ Lesson Plan Review")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xe8, 0x65, 0x1a)

    sub = doc.add_paragraph("Generated by PlanCheck · K–12 Lesson Review Agent")
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x9c, 0x94, 0x85)

    def classify_row(row):
        bullets = row.get("section_feedback") or ["No issues found"]
        if any(any(w in b.lower() for w in ["missing","not found","absent"]) for b in bullets):
            return "red"
        if any(b.lower() != "no issues found" for b in bullets):
            return "orange"
        return "green"

    counts = {"green":0,"orange":0,"red":0}
    for row in rows: counts[classify_row(row)] += 1

    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(10)
    legend.paragraph_format.space_after = Pt(10)
    for i, (label, count, color) in enumerate([
        ("✓ Ready to go", counts["green"], RGBColor(0x2d,0x7a,0x4f)),
        ("⚠ Needs fine-tuning", counts["orange"], RGBColor(0xe8,0x65,0x1a)),
        ("✗ Failed review", counts["red"], RGBColor(0xc0,0x39,0x2b)),
    ]):
        if i > 0:
            sep = legend.add_run("   |   ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor(0xcc,0xcc,0xcc)
        r = legend.add_run(f"{label}: {count}")
        r.font.size = Pt(9); r.bold = True; r.font.color.rgb = color

    doc.add_paragraph()

    # Include decisions column if provided
    has_decisions = any(d.get("decision") not in (None, "pending") for d in decisions)
    col_count = 4 if has_decisions else 3
    col_widths = [Inches(1.4), Inches(2.6), Inches(2.6), Inches(1.0)] if has_decisions else [Inches(1.6), Inches(3.0), Inches(3.0)]

    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    headers = ["Section", "Formatting Feedback", "Section Feedback"]
    if has_decisions: headers.append("Decision")

    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"),"e8651a"); tcPr.append(shd)

    for idx, row in enumerate(rows):
        tr = table.add_row()
        cells = tr.cells
        cells[0].text = row.get("section","")
        cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        cells[0].paragraphs[0].runs[0].bold = True

        fill = "fff8f3" if idx % 2 == 0 else "ffffff"
        for cell in cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),fill); tcPr.append(shd)

        for col_idx, key in enumerate(["formatting_feedback","section_feedback"], start=1):
            bullets = row.get(key) or ["No issues found"]
            cell = cells[col_idx]; cell.text = ""
            for i, bullet in enumerate(bullets):
                p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.12)
                r = p.add_run(f"• {bullet}")
                r.font.size = Pt(8)
                b_lower = bullet.lower()
                if col_idx == 2:
                    if b_lower == "no issues found":
                        r.font.color.rgb = RGBColor(0x2d,0x7a,0x4f)
                    elif any(w in b_lower for w in ["missing","not found","absent"]):
                        r.font.color.rgb = RGBColor(0xc0,0x39,0x2b)
                    else:
                        r.font.color.rgb = RGBColor(0xe8,0x65,0x1a)
                else:
                    r.font.color.rgb = RGBColor(0x1a,0x1a,0x14)

        if has_decisions and len(decisions) > idx:
            d = decisions[idx]
            decision = d.get("decision","pending")
            note = d.get("override_note","")
            cell = cells[3]; cell.text = ""
            p = cell.paragraphs[0]
            label = {"accept":"✓ Accepted","override":"✏ Override","pending":"— Pending"}.get(decision,"—")
            color = {"accept":RGBColor(0x2d,0x7a,0x4f),"override":RGBColor(0xe8,0x65,0x1a)}.get(decision, RGBColor(0x9c,0x94,0x85))
            r = p.add_run(label); r.font.size=Pt(8); r.bold=True; r.font.color.rgb=color
            if note:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(note); r2.font.size=Pt(7)
                r2.font.color.rgb=RGBColor(0x5c,0x57,0x48)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf


@app.route("/download-docx", methods=["POST"])
def download_docx():
    data = request.get_json()
    rows = data.get("rows", [])
    decisions = data.get("decisions", [])
    buf = build_review_docx(rows, decisions)
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="lesson-review.docx")


@app.route("/generate", methods=["POST"])
def generate():
    """Generate a fully written, standards-compliant lesson plan using Claude."""
    import anthropic
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = request.get_json()
    rows = data.get("rows", [])
    decisions = data.get("decisions", [])

    # ── Build prompt for Claude to write the full lesson plan ──────────────────
    sections_context = []
    for idx, row in enumerate(rows):
        d = decisions[idx] if idx < len(decisions) else {}
        decision = d.get("decision", "pending")
        override_note = d.get("override_note", "")
        issues = [b for b in (row.get("section_feedback") or []) + (row.get("formatting_feedback") or [])
                  if b.lower() != "no issues found"]

        if decision == "accept":
            sections_context.append(
                f"SECTION: {row['section']}\nSTATUS: Accepted as-is. Reproduce the content faithfully, applying only formatting corrections (fix spelling, abbreviations, formatting to standard)."
            )
        elif decision == "override" and override_note:
            issues_str = "; ".join(issues) if issues else "none"
            sections_context.append(
                f"SECTION: {row['section']}\nSTATUS: Override. Rewrite this section based on the teacher instruction: \"{override_note}\"\nOriginal issues to fix: {issues_str}"
            )
        else:
            sections_context.append(
                f"SECTION: {row['section']}\nSTATUS: Pending — include as-is with issues noted inline."
            )

    prompt = """You are an expert K-12 curriculum writer. Generate a complete, fully written, standards-compliant lesson plan document based on the section instructions below.

FORMATTING STANDARDS TO ENFORCE THROUGHOUT:
- Grade level format: Grade X (e.g. Grade 5)
- Duration format: X min (e.g. 45 min)
- Learning objectives must start with Bloom's Taxonomy action verbs (identify, explain, analyze, evaluate, create, compare, demonstrate)
- Standards must include both code AND description
- No weak verbs (understand, know, learn, appreciate)
- Fix all spelling errors, duplicate words, and formatting issues

For each section, output a JSON object in this array. Return ONLY valid JSON, no markdown.

[
  {
    "section": "<section name>",
    "content_type": "paragraphs" | "bullets" | "table_rows",
    "content": "<full written content as a single string for paragraphs, or array of strings for bullets, or array of [label, value] pairs for table_rows>"
  }
]

SECTION INSTRUCTIONS:
""" + "\n\n".join(sections_context)

    client = anthropic.Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response.content[0].text.strip()
    import re
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    sections = json.loads(raw)

    # ── Build the Word document ────────────────────────────────────────────────
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Formatted Lesson Plan")
    r.bold = True; r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0xe8, 0x65, 0x1a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Generated by PlanCheck · Standards-Compliant")
    rs.font.size = Pt(9); rs.font.color.rgb = RGBColor(0x9c, 0x94, 0x85)

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_after = Pt(12)
    div.border = None

    for sec_data in sections:
        section_name = sec_data.get("section", "")
        content_type = sec_data.get("content_type", "paragraphs")
        content = sec_data.get("content", "")

        # Section heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.border = {"bottom": {"style": "single", "size": 4, "color": "E8E2D6"}}
        hr = h.add_run(section_name)
        hr.bold = True; hr.font.size = Pt(13)
        hr.font.color.rgb = RGBColor(0x2d, 0x7a, 0x4f)

        if content_type == "bullets" and isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
                p.add_run(str(item)).font.size = Pt(11)

        elif content_type == "table_rows" and isinstance(content, list):
            if content:
                border = {"style": "single", "size": 4, "color": "CCCCCC"}
                tbl = doc.add_table(rows=1, cols=2)
                tbl.style = "Table Grid"
                # Remove header row, fill data rows
                tbl.rows[0].cells[0].text = ""
                tbl.rows[0].cells[1].text = ""
                for row_data in content:
                    if isinstance(row_data, list) and len(row_data) >= 2:
                        row = tbl.add_row()
                        lc = row.cells[0]
                        vc = row.cells[1]
                        lc.width = Inches(1.8)
                        vc.width = Inches(5.0)
                        lc.paragraphs[0].clear()
                        lr = lc.paragraphs[0].add_run(str(row_data[0]))
                        lr.bold = True; lr.font.size = Pt(10)
                        lr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x14)
                        vc.paragraphs[0].clear()
                        vr = vc.paragraphs[0].add_run(str(row_data[1]))
                        vr.font.size = Pt(10)
                # Remove the blank header row
                tbl._tbl.remove(tbl.rows[0]._tr)

        else:
            # Paragraphs
            text = content if isinstance(content, str) else " ".join(content)
            for para_text in text.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(para_text.strip()).font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="formatted-lesson-plan.docx")


@app.route("/create", methods=["POST"])
def create():
    """Generate a full sub lesson plan from teacher answers."""
    import anthropic, re
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.get_json()
    grade     = data.get("grade", "")
    subject   = data.get("subject", "")
    duration  = data.get("duration", "45 min")
    days      = data.get("days", "1 day")
    unit      = data.get("unit", "")
    topic     = data.get("topic", "")
    level     = data.get("level", "")
    materials = data.get("materials", "")
    notes     = data.get("notes", "")

    prompt = f"""You are an expert K-12 curriculum writer. A teacher is going to be absent and needs a complete, ready-to-use lesson plan for their substitute teacher.

TEACHER INPUTS:
- Grade: {grade}
- Subject: {subject}
- Class duration: {duration}
- Days absent: {days}
- Current unit and what students learned last week: {unit}
- What the sub should cover: {topic}
- Sub commitment level: {level}
- Where to find materials: {materials}
- Additional notes: {notes}

Write a complete, formatted substitute lesson plan. The sub should be able to pick this up and run the class with zero preparation.

Return ONLY a JSON array with no markdown fences. Each element is one section:
[
  {{
    "section": "<section name>",
    "content_type": "paragraphs" | "bullets" | "table_rows",
    "content": "<string for paragraphs> | <array of strings for bullets> | <array of [label, value] for table_rows>"
  }}
]

Include these sections in order:
1. Lesson Title - a clear, specific title for the lesson
2. Grade Level and Subject - table_rows with Grade and Subject
3. Duration - table_rows with Duration and Days
4. Learning Objectives - bullets, 2-4 measurable objectives using Bloom verbs, written for the sub to read aloud or post
5. Materials and Where to Find Them - bullets listing each material and exactly where it is
6. Schedule and Timing - table_rows with time blocks and what happens in each (e.g. 0-5 min: Settle in and attendance)
7. Step-by-Step Instructions - paragraphs with numbered, very explicit instructions the sub follows in order
8. What Students Should Produce - paragraphs describing the expected output or completed work
9. If Students Finish Early - bullets with approved extension activities
10. Classroom Routines - bullets covering routines the sub needs to know (attendance, bathroom, transitions)
11. Students Who Need Extra Support - paragraphs describing any students needing attention (use first names only)
12. Emergency and Contact Info - table_rows with key contacts and procedures

Write in plain, direct language for the sub. Be specific and leave nothing ambiguous. Do not use weak language like "feel free to" or "you may want to". Give clear instructions."""

    client = anthropic.Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    sections = json.loads(raw)

    # Build Word doc
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Header banner
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Substitute Lesson Plan")
    r.bold = True; r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0xe8, 0x65, 0x1a)

    sub_note = doc.add_paragraph()
    sub_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = sub_note.add_run(f"Generated by RainCheck  |  {grade}  |  {subject}")
    rn.font.size = Pt(10); rn.font.color.rgb = RGBColor(0x9c, 0x94, 0x85)

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rnotice = notice.add_run("Please read this plan fully before class begins.")
    rnotice.bold = True; rnotice.font.size = Pt(10)
    rnotice.font.color.rgb = RGBColor(0x2d, 0x7a, 0x4f)
    doc.add_paragraph()

    for sec_data in sections:
        section_name = sec_data.get("section", "")
        content_type = sec_data.get("content_type", "paragraphs")
        content = sec_data.get("content", "")

        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        hr = h.add_run(section_name)
        hr.bold = True; hr.font.size = Pt(13)
        hr.font.color.rgb = RGBColor(0x2d, 0x7a, 0x4f)

        if content_type == "bullets" and isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
                p.add_run(str(item)).font.size = Pt(11)

        elif content_type == "table_rows" and isinstance(content, list):
            if content:
                tbl = doc.add_table(rows=0, cols=2)
                tbl.style = "Table Grid"
                for row_data in content:
                    if isinstance(row_data, list) and len(row_data) >= 2:
                        row = tbl.add_row()
                        lc, vc = row.cells[0], row.cells[1]
                        lc.width = Inches(2.0); vc.width = Inches(4.8)
                        lr = lc.paragraphs[0].add_run(str(row_data[0]))
                        lr.bold = True; lr.font.size = Pt(10)
                        lr.font.color.rgb = RGBColor(0xe8, 0x65, 0x1a)
                        vr = vc.paragraphs[0].add_run(str(row_data[1]))
                        vr.font.size = Pt(10)

        else:
            text = content if isinstance(content, str) else " ".join(str(x) for x in content)
            for para_text in text.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(para_text.strip()).font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="sub-lesson-plan.docx")


if __name__ == "__main__":
    print("Running at http://127.0.0.1:5001")
    app.run(debug=False, host="0.0.0.0", port=int(os.environ.get("PORT", 5001)))