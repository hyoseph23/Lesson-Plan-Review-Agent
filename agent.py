"""
K-12 Lesson Plan & Worksheet Review Agent
Ingests a .docx, .pdf, or .txt document and returns a structured review table.
"""

import re
import json
from pathlib import Path
from docx import Document
import anthropic

# ── Section order (matches K-12 lesson plan template) ─────────────────────────
SECTION_ORDER = [
    "1. Header & Logo",
    "2. Lesson Title",
    "3. Grade Level & Subject",
    "4. Duration",
    "5. Learning Objectives",
    "6. Standards Alignment",
    "7. Materials & Resources",
    "8A. Introduction / Hook",
    "8B. Direct Instruction",
    "8C. Guided Practice",
    "8D. Independent Practice",
    "8E. Closure",
    "9. Differentiation",
    "10. Assessment",
    "11. Teacher Notes",
]

# ── Formatting rules (deterministic, pre-LLM pass) ────────────────────────────
FORMATTING_RULES = {
    "double_space": {
        "pattern": r"\.  +",
        "message": "Use one space after a period.",
    },
    "duplicate_words": {
        "pattern": r"\b(\w+)\s+\1\b",
        "message": "Duplicate word found back-to-back.",
    },
    "all_caps_words": {
        "pattern": r"\b[A-Z]{4,}\b",
        "message": "Avoid unnecessary ALL CAPS — use bold or italics for emphasis instead.",
    },
    "grade_format": {
        "pattern": r"\b\d+th\s+grade\b|\b\d+st\s+grade\b|\b\d+nd\s+grade\b|\b\d+rd\s+grade\b",
        "message": "Grade level should follow format: Grade X (e.g. Grade 5, not 5th grade).",
    },
    "time_format": {
        "pattern": r"\b\d+\s*minutes\b",
        "message": "Duration should follow format: X min (e.g. 45 min, not 45 minutes).",
    },
    "weak_objective_verb": {
        "pattern": r"\b(understand|know|learn|appreciate|be aware of)\b",
        "message": 'Weak or unmeasurable verb — replace with a Bloom\'s Taxonomy action verb (e.g. identify, explain, analyze, create).',
    },
}


def check_header_logo(file_path: str) -> bool:
    """Return True if the docx header contains at least one image (school logo)."""
    ext = file_path.rsplit(".", 1)[-1].lower()
    if ext != "docx":
        return False
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(file_path) as z:
            names = z.namelist()
            header_files = [n for n in names if n.startswith("word/header")]
            for hf in header_files:
                xml = z.read(hf)
                root = ET.fromstring(xml)
                ns = {
                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }
                if root.findall(".//w:drawing", ns):
                    return True
                if root.findall(".//{urn:schemas-microsoft-com:vml}shape"):
                    return True
    except Exception:
        pass
    return False


def extract_document_text(file_path: str) -> dict:
    """Extract full text and table contents from docx, txt, or pdf."""
    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "txt":
        text = Path(file_path).read_text(errors="ignore")
        return {"paragraphs": text.splitlines(), "tables": [], "full_text": text}

    if ext == "pdf":
        try:
            import subprocess
            result = subprocess.run(["pdftotext", file_path, "-"], capture_output=True, text=True)
            text = result.stdout
        except Exception:
            text = ""
        return {"paragraphs": text.splitlines(), "tables": [], "full_text": text}

    # Default: docx
    doc = Document(file_path)
    full_paragraphs = []
    tables_content = []
    all_text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_paragraphs.append(para.text.strip())
            all_text_parts.append(para.text.strip())

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
                if cell_text:
                    all_text_parts.append(cell_text)
            table_rows.append(row_cells)
        tables_content.append(table_rows)

    return {
        "paragraphs": full_paragraphs,
        "tables": tables_content,
        "full_text": "\n".join(all_text_parts),
    }


def run_formatting_checks(text: str) -> list[dict]:
    """Run deterministic regex-based formatting checks on full document text."""
    issues = []

    for rule_name, rule in FORMATTING_RULES.items():
        matches = list(re.finditer(rule["pattern"], text, re.IGNORECASE))

        # Filter duplicate words — skip trivial false positives
        if rule_name == "duplicate_words":
            matches = [
                m for m in matches
                if len(m.group(1)) > 2 and m.group(1).lower() not in {"the", "a", "to", "of", "in"}
            ]

        # Filter all-caps — skip known education acronyms
        if rule_name == "all_caps_words":
            known_acronyms = {"ELL", "IEP", "ESL", "STEM", "STEAM", "ELA", "PDF", "URL",
                              "DOK", "UDL", "PBL", "NGSS", "CCSS", "SEL"}
            matches = [m for m in matches if m.group(0) not in known_acronyms]

        for match in matches:
            issues.append({
                "rule": rule_name,
                "found": match.group(0),
                "message": rule["message"],
                "position": match.start(),
            })

    return issues


def build_llm_prompt(doc_data: dict, formatting_issues: list[dict]) -> str:
    """Construct the prompt for the LLM review pass."""
    formatting_summary = ""
    if formatting_issues:
        by_rule = {}
        for issue in formatting_issues:
            by_rule.setdefault(issue["rule"], []).append(f'"{issue["found"]}"')
        lines = []
        for rule, found_list in by_rule.items():
            lines.append(f"- {FORMATTING_RULES[rule]['message']} Found: {', '.join(found_list[:5])}")
        formatting_summary = "\n".join(lines)
    else:
        formatting_summary = "None detected by pre-scan."

    tables_text = ""
    for i, table in enumerate(doc_data["tables"]):
        tables_text += f"\n[Table {i+1}]\n"
        for row in table:
            tables_text += " | ".join(row) + "\n"

    prompt = f"""You are an expert K-12 curriculum reviewer. Your job is to review a teacher's lesson plan or worksheet and return structured, specific feedback.

## DOCUMENT TEXT
{doc_data['full_text']}

## TABLES IN DOCUMENT
{tables_text}

## PRE-SCANNED FORMATTING ISSUES (already detected)
{formatting_summary}

## TASK
Review the document against the rules below and return a JSON array — one element per section, in order. Do not skip any section; if missing, still include it and flag it.

## SECTION ORDER (use these exact names)
2. Lesson Title
3. Grade Level & Subject
4. Duration
5. Learning Objectives
6. Standards Alignment
7. Materials & Resources
8A. Introduction / Hook
8B. Direct Instruction
8C. Guided Practice
8D. Independent Practice
8E. Closure
9. Differentiation
10. Assessment
11. Teacher Notes

## FORMATTING & GRAMMAR RULES (apply document-wide, note per section where they appear)
- No misspellings — flag each with the word found
- No duplicate words back-to-back
- Avoid unnecessary ALL CAPS — use bold or italics for emphasis
- Grade level must follow: Grade X (e.g. Grade 5, not "5th grade")
- Duration must follow: X min (e.g. 45 min, not "45 minutes")
- Weak/unmeasurable verbs in objectives (understand, know, learn, appreciate) must be replaced with Bloom's Taxonomy action verbs
- One space after a period

## SECTION-SPECIFIC RULES

**2. Lesson Title**
- Must be clear and descriptive — flag if vague or generic (e.g. "Science Lesson", "Math Worksheet")
- Must appear at the top of the document
- Capitalize first letter of each word only — flag ALL CAPS titles

**3. Grade Level & Subject**
- Must clearly state grade level (format: Grade X) AND subject area
- Flag if either is missing
- Flag if grade level format is wrong (e.g. "5th grade" instead of "Grade 5")

**4. Duration**
- Must state total lesson duration in format: X min (e.g. 45 min)
- Flag if missing
- Flag if duration seems unrealistic for the grade level (e.g. 180 min for Grade 2)

**5. Learning Objectives**
- Must include 2–5 clear, measurable objectives
- Each must start with a Bloom's Taxonomy action verb (identify, explain, analyze, evaluate, create, compare, demonstrate, classify, construct, design)
- Flag weak verbs: understand, know, learn, appreciate, be aware of
- Flag if fewer than 2 or more than 5 objectives
- Each objective should be written as: "Students will be able to [verb] [what]"
- Flag if any objective exceeds 2 lines

**6. Standards Alignment**
- Must reference at least one standard (Common Core, NGSS, or state standard)
- Must include both the standard code AND a brief description
- Flag if missing entirely
- Flag if only a code is listed with no description

**7. Materials & Resources**
- Must list all materials and resources needed
- Flag if missing
- Flag if vague (e.g. "various supplies" without specifics)
- Flag if digital resources are listed without titles or URLs

**8A. Introduction / Hook**
- Must include an engaging opening to activate prior knowledge
- Flag if missing
- Flag if over 5 lines
- Flag if no clear connection to learning objectives

**8B. Direct Instruction**
- Must describe what the teacher will teach and how
- Must mention key vocabulary or concepts being introduced
- Flag if missing
- Flag if no mention of checking for understanding
- Flag if over 10 lines without clear sub-steps

**8C. Guided Practice**
- Must describe a structured activity with teacher support
- Flag if missing
- Flag if instructions are unclear
- Flag if over 8 lines

**8D. Independent Practice**
- Must describe an activity students complete on their own
- Flag if missing
- Flag if no clear success criteria or instructions
- Flag if over 8 lines

**8E. Closure**
- Must include a specific closing activity (exit ticket, reflection, class discussion)
- Flag if missing
- Flag if closure is generic (e.g. just "review the lesson")
- Flag if over 5 lines

**9. Differentiation**
- Must address at least two groups: ELL students, students with IEPs, advanced learners, or struggling learners
- Flag if missing entirely
- Flag if only one group is addressed
- Flag if strategies are too generic (e.g. "provide extra help" or "extend the activity")

**10. Assessment**
- Must describe how student learning will be measured
- Must connect to at least one learning objective
- Flag if missing
- Flag if only "participation" is listed
- Flag if assessment does not align with any stated objective

**11. Teacher Notes**
- Optional — do NOT flag if missing
- Flag if notes contain content that belongs in another section
- Flag if over 5 lines

## OUTPUT FORMAT
Return ONLY a valid JSON array. No markdown, no explanation. Each element:
{{
  "section": "<section name exactly as listed above>",
  "formatting_feedback": ["<bullet 1>", "<bullet 2>"],
  "section_feedback": ["<bullet 1>", "<bullet 2>"]
}}

If a column has no issues use: ["No issues found"]
Be specific — name the exact problem and where it appears.
"""
    return prompt


# Canonical numbered section names (header row handled separately)
CANONICAL_SECTIONS = [
    "2. Lesson Title",
    "3. Grade Level & Subject",
    "4. Duration",
    "5. Learning Objectives",
    "6. Standards Alignment",
    "7. Materials & Resources",
    "8A. Introduction / Hook",
    "8B. Direct Instruction",
    "8C. Guided Practice",
    "8D. Independent Practice",
    "8E. Closure",
    "9. Differentiation",
    "10. Assessment",
    "11. Teacher Notes",
]


def run_llm_review(prompt: str) -> list[dict]:
    """Send prompt to Claude and parse the JSON response."""
    client = anthropic.Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    rows = json.loads(raw)

    # Strip any header/logo rows Claude may have added
    rows = [r for r in rows if not any(
        w in r.get("section", "").lower() for w in ["header", "logo", "branding"]
    )]

    # Remap by position onto canonical names
    remapped = []
    for i, row in enumerate(rows):
        canonical_name = CANONICAL_SECTIONS[i] if i < len(CANONICAL_SECTIONS) else row.get("section", f"Section {i+2}")
        remapped.append({
            "section": canonical_name,
            "formatting_feedback": row.get("formatting_feedback", ["No issues found"]),
            "section_feedback": row.get("section_feedback", ["No issues found"]),
        })

    # Pad missing sections
    for i in range(len(remapped), len(CANONICAL_SECTIONS)):
        remapped.append({
            "section": CANONICAL_SECTIONS[i],
            "formatting_feedback": ["No issues found"],
            "section_feedback": ["Section was not found — may be missing from document."],
        })

    return remapped


def format_table_terminal(rows: list[dict]) -> str:
    """Print a readable table to terminal."""
    col1_w = max(len(r["section"]) for r in rows) + 2
    col2_w = 50
    col3_w = 50

    def wrap(text, width):
        words = text.split()
        lines, line = [], []
        for w in words:
            if sum(len(x) + 1 for x in line) + len(w) > width:
                lines.append(" ".join(line))
                line = [w]
            else:
                line.append(w)
        if line:
            lines.append(" ".join(line))
        return lines

    sep = f"+{'-'*col1_w}+{'-'*col2_w}+{'-'*col3_w}+"
    header = f"|{'Document Section'.center(col1_w)}|{'Formatting & Grammar Feedback'.center(col2_w)}|{'Section Feedback'.center(col3_w)}|"
    output = [sep, header, sep]

    for row in rows:
        fmt_bullets = row.get("formatting_feedback", ["No issues found"])
        sec_bullets = row.get("section_feedback", ["No issues found"])
        fmt_lines = []
        for b in fmt_bullets:
            fmt_lines.extend(wrap(f"• {b}", col2_w - 2))
        sec_lines = []
        for b in sec_bullets:
            sec_lines.extend(wrap(f"• {b}", col3_w - 2))
        n_lines = max(1, len(fmt_lines), len(sec_lines))
        section_lines = wrap(row["section"], col1_w - 2)
        section_lines += [""] * (n_lines - len(section_lines))
        fmt_lines += [""] * (n_lines - len(fmt_lines))
        sec_lines += [""] * (n_lines - len(sec_lines))
        for i in range(n_lines):
            c1 = section_lines[i] if i < len(section_lines) else ""
            c2 = fmt_lines[i] if i < len(fmt_lines) else ""
            c3 = sec_lines[i] if i < len(sec_lines) else ""
            output.append(f"|{c1:<{col1_w}}|{c2:<{col2_w}}|{c3:<{col3_w}}|")
        output.append(sep)

    return "\n".join(output)


def review_document(file_path: str, output_json: str = None) -> list[dict]:
    """Main entry point. Returns list of review row dicts."""
    print(f" Reading document: {file_path}")
    doc_data = extract_document_text(file_path)

    print("🔍 Checking header for school logo...")
    has_logo = check_header_logo(file_path)
    header_row = {
        "section": "1. Header & Branding",
        "formatting_feedback": ["No issues found"],
        "section_feedback": ["No issues found"] if has_logo else [
            "School logo or branding is missing from the document header. Add your school logo to the top of the document."
        ],
    }
    print(f"   Logo found: {has_logo}")

    print("Running formatting pre-scan...")
    formatting_issues = run_formatting_checks(doc_data["full_text"])
    print(f"   Found {len(formatting_issues)} formatting issue(s)")

    print("Sending to the agent for full section review...")
    prompt = build_llm_prompt(doc_data, formatting_issues)
    rows = run_llm_review(prompt)
    print(f"   Received {len(rows)} section reviews")

    rows = [header_row] + rows

    if output_json:
        Path(output_json).write_text(json.dumps(rows, indent=2))
        print(f" JSON saved to {output_json}")

    return rows


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python agent.py <path/to/lesson_plan.docx> [output.json]")
        sys.exit(1)

    file_path = sys.argv[1]
    output_json = sys.argv[2] if len(sys.argv) > 2 else None

    rows = review_document(file_path, output_json)
    print("\n" + format_table_terminal(rows))